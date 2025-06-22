import os
import asyncio
import openai
from datetime import datetime, timedelta
from openai import AsyncOpenAI
from uuid import uuid4
from fastapi import FastAPI, Request, Response, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse

# Load environment variables (API key, model, etc.)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

openai.api_key = os.getenv("OPENAI_API_KEY")
MODEL_NAME = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
PROMPT_PATH = os.getenv("PROMPT_PATH")
TEMPLATE_PATH = os.getenv("TEMPLATE_PATH")
SYSTEM_PROMPT = None
if PROMPT_PATH:
    try:
        with open(PROMPT_PATH, 'r') as f:
            SYSTEM_PROMPT = f.read().strip()
    except Exception as e:
        print(f"Warning: Could not load system prompt from {PROMPT_PATH}: {e}")

app = FastAPI()

# Mount static files directory
app.mount("/static", StaticFiles(directory="static"), name="static")

# Set up templates directory
templates = Jinja2Templates(directory="templates")

# In-memory storage for chat sessions and generated docs
sessions = {}     # { session_id: {"messages": [...], "questions_asked": int, "doc_ready": bool, "last_active": datetime} }
session_docs = {}  # { session_id: bytes of generated .docx document }

# Helper function to create a Word document from text (runs in a thread to avoid blocking)
from io import BytesIO
from docx import Document
def create_doc_from_text(text: str) -> bytes:
    if TEMPLATE_PATH and os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
    # Add content to the Word document (could be formatted as needed)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    """Serve the main index page with a form for initial inputs."""
    # Ensure session cookie exists
    cleanup_sessions()
    session_id = request.cookies.get("session_id")
    if not session_id:
        session_id = str(uuid4())
    # Render the index.html template
    response = templates.TemplateResponse("index.html", {"request": request})
    # Set session cookie (HttpOnly since we only need it on server side)
    response.set_cookie(key="session_id", value=session_id, httponly=True)
    return response

@app.get("/rams_chat", response_class=HTMLResponse)
async def rams_chat(request: Request):
    """Serve the chat interface page."""
    cleanup_sessions()
    session_id = request.cookies.get("session_id")
    if not session_id:
        session_id = str(uuid4())
    # Initialize session if new
    if session_id not in sessions:
        sessions[session_id] = {"messages": [], "questions_asked": 0, "doc_ready": False}
    # Update last active time
    sessions[session_id]["last_active"] = datetime.now()
    # Add an initial assistant prompt if conversation is empty
    if len(sessions[session_id]["messages"]) == 0:
        initial_prompt = "Please describe the task you'd like a RAMS for. Include as much detail as possible."
        sessions[session_id]["messages"].append({"role": "assistant", "content": initial_prompt})
    # Render the chat page with current conversation
    response = templates.TemplateResponse("rams_chat.html", {
        "request": request,
        "messages": sessions[session_id]["messages"]
    })
    response.set_cookie(key="session_id", value=session_id, httponly=True)
    return response

@app.post("/rams_chat", response_class=HTMLResponse)
async def start_chat(request: Request):
    """Handle form submission from index page and redirect to chat interface."""
    cleanup_sessions()
    form_data = await request.form()
    session_id = request.cookies.get("session_id") or str(uuid4())
    # Initialize a new session for these answers
    sessions[session_id] = {"messages": [], "questions_asked": 0, "doc_ready": False}
    # Combine all provided answers into one user message (for context)
    answers = []
    for i in range(1, 8):
        ans = form_data.get(f"answer{i}")
        if ans:
            answers.append(f"Answer {i}: {ans}")
    if answers:
        combined_answers = "\n".join(answers)
        sessions[session_id]["messages"].append({"role": "user", "content": combined_answers})
        try:
            messages_for_openai = []
            if sessions[session_id]["questions_asked"] == 0:
                if SYSTEM_PROMPT:
                    messages_for_openai.append({"role": "system", "content": SYSTEM_PROMPT})
                system_instruction = (
                    "You are an expert safety engineer. Ask the user one follow-up question needed "
                    "to create a detailed Risk Assessment Method Statement (RAMS). "
                    "Keep your question concise."
                )
                messages_for_openai.append({"role": "system", "content": system_instruction})
            messages_for_openai += sessions[session_id]["messages"]
            openai_resp = await client.chat.completions.create(
                model=MODEL_NAME,
                messages=messages_for_openai,
                temperature=0.7
            )
            first_question = openai_resp.choices[0].message.content.strip()
        except Exception as e:
            first_question = f"(Error generating question: {e})"
        sessions[session_id]["messages"].append({"role": "assistant", "content": first_question})
        sessions[session_id]["questions_asked"] += 1
    # Render the chat interface with the provided answers in context (now including the first AI question if generated)
    response = templates.TemplateResponse("rams_chat.html", {
        "request": request,
        "messages": sessions[session_id]["messages"]
    })
    # Set/update session cookie
    response.set_cookie(key="session_id", value=session_id, httponly=True)
    # Update last active time for the session
    sessions[session_id]["last_active"] = datetime.now()
    return response

@app.post("/chat", response_class=JSONResponse)
async def chat_api(request: Request):
    """Endpoint for AJAX calls from the chat interface (sending a new message and getting a response)."""
    cleanup_sessions()
    data = await request.json()
    user_message = data.get("message", "").strip()
    session_id = request.cookies.get("session_id")
    if not session_id:
        # Create a new session if no cookie was present (edge case)
        session_id = str(uuid4())
        sessions[session_id] = {"messages": [], "questions_asked": 0, "doc_ready": False}
    # Ensure session exists in memory
    if session_id not in sessions:
        sessions[session_id] = {"messages": [], "questions_asked": 0, "doc_ready": False}
    session_data = sessions[session_id]
    # Append the user's message to conversation history
    if user_message:
        session_data["messages"].append({"role": "user", "content": user_message})
    # If a final document has already been generated for this session, inform the user
    if session_data.get("doc_ready"):
        sessions[session_id]["last_active"] = datetime.now()
        return {"answer": "The RAMS document has already been generated.", "done": True}
    # Determine whether to ask another question or generate the final document
    if session_data["questions_asked"] < 20:
        # Use OpenAI to generate the next follow-up question
        try:
            # Include a guiding system instruction on the first question only
            messages_for_openai = []
            if session_data["questions_asked"] == 0:
                if SYSTEM_PROMPT:
                    messages_for_openai.append({"role": "system", "content": SYSTEM_PROMPT})
                system_instruction = (
                    "You are an expert safety engineer. Ask the user one follow-up question needed "
                    "to create a detailed Risk Assessment Method Statement (RAMS). "
                    "Keep your question concise."
                )
                messages_for_openai.append({"role": "system", "content": system_instruction})
            # Provide the conversation history (including any initial user inputs)
            messages_for_openai += session_data["messages"]
            # Call the OpenAI ChatCompletion endpoint asynchronously
            openai_resp = await client.chat.completions.create(
                model=MODEL_NAME,
                messages=messages_for_openai,
                temperature=0.7
            )
            assistant_content = openai_resp["choices"][0]["message"]["content"].strip()
        except Exception as e:
            assistant_content = f"(Error generating question: {e})"
        # Append the AI's question to the history
        session_data["messages"].append({"role": "assistant", "content": assistant_content})
        session_data["questions_asked"] += 1
        # Update last active time
        sessions[session_id]["last_active"] = datetime.now()
        return {"answer": assistant_content}
    else:
        # After enough questions have been asked, use OpenAI to generate the final RAMS content
        try:
            finalize_instruction = "Now use the information gathered to produce the full detailed RAMS content for the task."
            messages_for_openai = session_data["messages"] + [{"role": "system", "content": finalize_instruction}]
            final_resp = await client.chat.completions.create(
                model=MODEL_NAME,
                messages=messages_for_openai,
                temperature=0.7
            )
            final_content = final_resp["choices"][0]["message"]["content"].strip()
        except Exception as e:
            final_content = f"(Error generating final content: {e})"
        # (Optionally, you might append the final content to messages, but we'll mark it as generated)
        session_data["messages"].append({"role": "assistant", "content": "[RAMS document generated]"})
        # Create the Word document in a background thread
        doc_bytes = await asyncio.to_thread(create_doc_from_text, final_content)
        session_docs[session_id] = doc_bytes
        session_data["doc_ready"] = True
        # Update last active time
        sessions[session_id]["last_active"] = datetime.now()
        return {"answer": "Your RAMS document is ready for download.", "done": True}

@app.get("/download")
async def download_doc(request: Request):
    """Provide the generated RAMS Word document as a downloadable file."""
    cleanup_sessions()
    session_id = request.cookies.get("session_id")
    if not session_id or session_id not in session_docs:
        raise HTTPException(status_code=404, detail="No document available for this session.")
    doc_data = session_docs[session_id]
    # Return the Word document as an attachment
    sessions[session_id]["last_active"] = datetime.now()
    return Response(doc_data, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=RAMS.docx"})


